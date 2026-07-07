"""
Agent 8 – Migration Runbook Generator
Generates a full migration runbook and exports to PDF, DOCX and Markdown.
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from tenacity import retry, stop_after_attempt, wait_exponential

from app.models.schemas import AgentState, MigrationRunbook, RiskLevel, RunbookSection
from app.prompts.runbook_prompt import RUNBOOK_SYSTEM_PROMPT, RUNBOOK_USER_TEMPLATE
from app.services import get_logger, get_llm
from config.settings import get_settings

log = get_logger(__name__)
settings = get_settings()

_OUTPUT = Path(settings.reports_output_dir)

# ── Rule-based default runbook sections ──────────────────────────────────────
_DEFAULT_SECTIONS = [
    RunbookSection(
        title="1. Executive Summary & Project Charter",
        order=1,
        content="High-level overview of the migration project, business justification, and scope.",
        tasks=[
            "Define project scope and boundaries",
            "Identify key stakeholders and executive sponsor",
            "Establish project governance structure",
            "Define success criteria and KPIs",
            "Create project charter document",
        ],
        estimated_duration="2 weeks",
        responsible_team="PMO & Executive Leadership",
    ),
    RunbookSection(
        title="2. Project Preparation & Landscape Assessment",
        order=2,
        content="Technical and functional preparation activities before conversion starts.",
        tasks=[
            "Run SAP Readiness Check tool",
            "Download and review Simplification List",
            "Execute Custom Code Migration Worklist (CCMW)",
            "Install SAP S/4HANA Migration Cockpit",
            "Set up sandbox conversion system",
            "Baseline system performance measurements",
        ],
        estimated_duration="4 weeks",
        responsible_team="SAP BASIS + ABAP Team",
    ),
    RunbookSection(
        title="3. Custom Code Adaptation",
        order=3,
        content="Remediation of all custom ABAP code flagged during ATC assessment.",
        tasks=[
            "Prioritize ATC findings by severity (Critical → High → Medium)",
            "Assign objects to ABAP developers",
            "Fix obsolete API usage (LIST_FROM_MEMORY, WRITE_FORM, etc.)",
            "Migrate ALV reports to CL_SALV_TABLE",
            "Fix SQL injection vulnerabilities",
            "Update deprecated function module calls",
            "Re-run ATC checks to verify fixes",
            "Perform unit testing for all changed objects",
            "Transport to quality system",
        ],
        estimated_duration="8 weeks",
        responsible_team="ABAP Development Team",
    ),
    RunbookSection(
        title="4. Business Partner Conversion",
        order=4,
        content="Mandatory conversion of Customer/Vendor master data to Business Partner concept.",
        tasks=[
            "Analyze current Customer and Vendor master data quality",
            "Run BP data quality checks (transaction BP)",
            "Execute pre-conversion validation reports",
            "Run transaction FLBP for BP assignment",
            "Verify BP roles: FLVN00 (Vendor) and FLCU00 (Customer)",
            "Test in sandbox system",
            "Document exceptions and manual corrections",
        ],
        estimated_duration="4 weeks",
        responsible_team="Finance Functional Team + ABAP",
    ),
    RunbookSection(
        title="5. Material Ledger Activation",
        order=5,
        content="Activate Material Ledger for all plants (mandatory in S/4HANA).",
        tasks=[
            "Identify all plants requiring ML activation",
            "Review current price control settings",
            "Execute ML activation in sandbox",
            "Validate inventory valuation",
            "Test material price changes",
            "Document impact on CO-PA and profitability analysis",
        ],
        estimated_duration="3 weeks",
        responsible_team="Finance & Controlling Team",
    ),
    RunbookSection(
        title="6. Universal Journal Migration",
        order=6,
        content="Migration to Universal Journal (ACDOCA) as single source of truth for finance.",
        tasks=[
            "Analyze FI/CO posting logic in custom programs",
            "Update custom code reading BSEG to use ACDOCA",
            "Migrate CO-PA configuration",
            "Test all financial reports",
            "Validate P&L and Balance Sheet output",
            "Reconcile totals between ECC and S/4HANA",
        ],
        estimated_duration="4 weeks",
        responsible_team="Finance Functional + ABAP Team",
    ),
    RunbookSection(
        title="7. Integration & Interface Testing",
        order=7,
        content="End-to-end testing of all interfaces, IDocs, and RFC connections.",
        tasks=[
            "Inventory all interfaces (RFC, IDoc, Web Services, APIs)",
            "Test RFC destinations to external systems",
            "Validate IDoc processing",
            "Test BAPIs and web service endpoints",
            "Regression test all custom function modules",
            "Verify PI/PO/Integration Suite connectivity",
        ],
        estimated_duration="4 weeks",
        responsible_team="Integration Team + BASIS",
    ),
    RunbookSection(
        title="8. User Acceptance Testing (UAT)",
        order=8,
        content="Business-driven acceptance testing covering all critical processes.",
        tasks=[
            "Prepare UAT test scripts from key business processes",
            "Set up UAT system with migrated data",
            "Execute P2P (Procure-to-Pay) scenarios",
            "Execute O2C (Order-to-Cash) scenarios",
            "Execute Finance month-end close simulation",
            "Document and resolve all defects",
            "Obtain sign-off from business process owners",
        ],
        estimated_duration="6 weeks",
        responsible_team="Business Process Owners + QA Team",
    ),
    RunbookSection(
        title="9. Cutover Planning & Execution",
        order=9,
        content="Detailed cutover plan ensuring minimal business downtime.",
        tasks=[
            "Define cutover weekend schedule",
            "Prepare cutover runbook with minute-by-minute tasks",
            "Execute mock cutover (minimum 2 rehearsals)",
            "Freeze source system and complete open transations",
            "Execute system conversion (SUM tool)",
            "Run post-conversion checks",
            "Release system for production use",
        ],
        estimated_duration="2 weeks",
        responsible_team="BASIS + PMO + All Teams",
    ),
    RunbookSection(
        title="10. Rollback Plan",
        order=10,
        content="Contingency plan if go-live must be aborted.",
        tasks=[
            "Define rollback decision criteria and go/no-go thresholds",
            "Take full system backup before cutover",
            "Document rollback procedure step by step",
            "Assign rollback team members and responsibilities",
            "Test rollback procedure in sandbox",
            "Define maximum tolerated downtime",
        ],
        estimated_duration="1 week (preparation)",
        responsible_team="BASIS + PMO",
    ),
    RunbookSection(
        title="11. Post Go-Live & Hypercare",
        order=11,
        content="Stabilization activities in the first weeks after go-live.",
        tasks=[
            "Establish hypercare support team (24/7 for first 2 weeks)",
            "Monitor system performance and workload",
            "Resolve P1/P2 incidents immediately",
            "Conduct daily business review calls",
            "Close out open ATC findings backlog",
            "Prepare lessons learned document",
            "Transition to BAU support",
        ],
        estimated_duration="4 weeks",
        responsible_team="Full Project Team",
    ),
]


@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=2, max=8))
def _llm_generate_runbook(state: AgentState, llm) -> list[RunbookSection]:
    key_risks = []
    if state.readiness_score:
        key_risks = state.readiness_score.critical_objects[:5]
    if state.simplification_report:
        key_risks += [i.title for i in state.simplification_report.business_partner_items[:2]]

    prompt = RUNBOOK_USER_TEMPLATE.format(
        system_id=state.sap_system.sid,
        readiness_score=state.readiness_score.overall_score if state.readiness_score else "N/A",
        risk_level=state.readiness_score.risk_level.value if state.readiness_score else "unknown",
        total_objects=state.custom_code.total_objects if state.custom_code else 0,
        total_findings=state.atc_report.total_findings if state.atc_report else 0,
        critical_count=state.atc_report.critical_count if state.atc_report else 0,
        effort_days=state.readiness_score.estimated_effort_days if state.readiness_score else 0,
        key_risks=", ".join(key_risks) or "None identified",
    )

    response = llm.invoke([
        {"role": "system", "content": RUNBOOK_SYSTEM_PROMPT},
        {"role": "user",   "content": prompt},
    ])
    content = response.content.strip()
    if "```json" in content:
        content = content.split("```json")[1].split("```")[0].strip()
    data = json.loads(content)
    sections = [RunbookSection(**s) for s in data.get("sections", [])]
    return sections


def _export_markdown(runbook: MigrationRunbook, path: Path) -> str:
    lines = [f"# {runbook.project_name}\n", f"**System:** {runbook.system_id}\n",
             f"**Generated:** {runbook.generated_at.strftime('%Y-%m-%d %H:%M UTC')}\n",
             f"**Estimated Duration:** {runbook.total_estimated_duration_weeks} weeks\n\n---\n"]
    for section in sorted(runbook.sections, key=lambda s: s.order):
        lines.append(f"\n## {section.title}\n")
        if section.estimated_duration:
            lines.append(f"**Duration:** {section.estimated_duration}  ")
        if section.responsible_team:
            lines.append(f"**Team:** {section.responsible_team}\n\n")
        lines.append(f"{section.content}\n\n")
        if section.tasks:
            lines.append("### Tasks\n")
            for task in section.tasks:
                lines.append(f"- [ ] {task}\n")
    output = "".join(lines)
    path.write_text(output, encoding="utf-8")
    return str(path)


def _export_pdf(runbook: MigrationRunbook, path: Path) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem

        doc = SimpleDocTemplate(str(path), pagesize=A4,
                                rightMargin=2*cm, leftMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle("Title", parent=styles["Title"],
                                     fontSize=20, spaceAfter=12, textColor=colors.HexColor("#1f2328"))
        story.append(Paragraph(runbook.project_name, title_style))
        story.append(Paragraph(f"System: {runbook.system_id} | Generated: {runbook.generated_at.strftime('%Y-%m-%d')}", styles["Normal"]))
        story.append(Paragraph(f"Estimated Project Duration: {runbook.total_estimated_duration_weeks} weeks", styles["Normal"]))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e5e7eb")))
        story.append(Spacer(1, 12))

        h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                                  fontSize=13, textColor=colors.HexColor("#3b82d4"), spaceAfter=6)
        for section in sorted(runbook.sections, key=lambda s: s.order):
            story.append(Paragraph(section.title, h2_style))
            if section.estimated_duration:
                story.append(Paragraph(f"<b>Duration:</b> {section.estimated_duration} | <b>Team:</b> {section.responsible_team or 'TBD'}", styles["Normal"]))
            story.append(Spacer(1, 4))
            story.append(Paragraph(section.content, styles["Normal"]))
            story.append(Spacer(1, 4))
            if section.tasks:
                items = [ListItem(Paragraph(f"☐ {t}", styles["Normal"])) for t in section.tasks]
                story.append(ListFlowable(items, bulletType="bullet", leftIndent=20))
            story.append(Spacer(1, 12))

        doc.build(story)
        return str(path)
    except Exception as exc:
        log.warning("PDF generation failed", error=str(exc))
        return ""


def _export_docx(runbook: MigrationRunbook, path: Path) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading(runbook.project_name, 0)
        doc.add_paragraph(f"System: {runbook.system_id} | Generated: {runbook.generated_at.strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Estimated Project Duration: {runbook.total_estimated_duration_weeks} weeks")
        doc.add_paragraph("─" * 60)

        for section in sorted(runbook.sections, key=lambda s: s.order):
            doc.add_heading(section.title, level=2)
            if section.estimated_duration:
                p = doc.add_paragraph()
                p.add_run(f"Duration: {section.estimated_duration}").bold = True
                p.add_run(f" | Team: {section.responsible_team or 'TBD'}")
            doc.add_paragraph(section.content)
            if section.tasks:
                doc.add_heading("Tasks:", level=3)
                for task in section.tasks:
                    doc.add_paragraph(f"☐ {task}", style="List Bullet")

        doc.save(str(path))
        return str(path)
    except Exception as exc:
        log.warning("DOCX generation failed", error=str(exc))
        return ""


def run_runbook_generation(state: AgentState) -> AgentState:
    """LangGraph node: generate the migration runbook."""
    log.info("Agent 8 – Runbook Generator started", assessment_id=state.assessment_id)
    try:
        _OUTPUT.mkdir(parents=True, exist_ok=True)

        # Try LLM-generated runbook, fall back to rule-based
        sections = _DEFAULT_SECTIONS
        try:
            llm = get_llm()
            use_llm = bool(llm.openai_api_key and "dummy" not in str(llm.openai_api_key))
            if use_llm:
                sections = _llm_generate_runbook(state, llm)
                log.info("Runbook generated via LLM")
        except Exception as exc:
            log.warning("LLM runbook generation failed, using rule-based", error=str(exc))

        effort_weeks = max(12, (state.readiness_score.estimated_effort_days // 5) if state.readiness_score else 20)

        runbook = MigrationRunbook(
            project_name=f"SAP ECC to S/4HANA Migration – {state.sap_system.sid}",
            system_id=state.sap_system.sid,
            sections=sections,
            total_estimated_duration_weeks=effort_weeks,
            generated_at=datetime.utcnow(),
        )

        aid = state.assessment_id[:8]
        md_path   = _OUTPUT / f"runbook_{aid}.md"
        pdf_path  = _OUTPUT / f"runbook_{aid}.pdf"
        docx_path = _OUTPUT / f"runbook_{aid}.docx"

        runbook.markdown_path = _export_markdown(runbook, md_path)
        runbook.pdf_path      = _export_pdf(runbook, pdf_path)
        runbook.docx_path     = _export_docx(runbook, docx_path)

        state.runbook = runbook
        state.steps_completed.append("runbook_generation")
        state.current_step = "dashboard_generation"
        log.info("Agent 8 – completed",
                 sections=len(sections),
                 duration_weeks=effort_weeks,
                 markdown_path=runbook.markdown_path)

    except Exception as exc:
        log.error("Agent 8 – failed", error=str(exc))
        state.error_messages.append(f"Runbook Generation: {exc}")

    return state
